import sys
sys.path.append('.../.../')
import tkinter as tk
import tkinter.ttk as ttk
from tkinter import ttk
from city.dom import dom_plosh
from city.dom import dom_energy
from city.comnata import  comnata_plosh
from city.comnata import  comnata_energy
from city.kvartura import kvartura_energy
from city.kvartura import  kvartura_plosh
import tkinter.messagebox as mb
import os
import docx

doc = docx.Document()
# Это кнопка для расчёта площади

def utog():
    pop1 = shirina.get()
    pop2 = dlina.get()
    pop3 = kolvo_com.get()
    if (pop1.isdigit()) and (pop2.isdigit()):
        vivod_ploshadi.delete(0,'end')
        y = comboExample.current()
        if y == 0:
            a = int(shirina.get())
            b = int(dlina.get())
            c = str(comnata_plosh(a, b))
            vivod_ploshadi.insert(0, c)
            c = ('Площадь комнаты ', str(comnata_plosh(a, b)), ' м^2')
            doc.add_paragraph(c)
        elif y == 1 and (len(pop3) != 0):
            d = int(kolvo_com.get())
            a = int(shirina.get())
            b = int(dlina.get())
            c = str(kvartura_plosh(a, b, d))
            vivod_ploshadi.insert(0, c)
            c = ('Площадь квартиры ', str(kvartura_plosh(a, b, d)), ' м^2')
            doc.add_paragraph(c)
        elif y == 2:
            a = int(shirina.get())
            b = int(dlina.get())
            c = str(dom_plosh(a, b))
            vivod_ploshadi.insert(0, c)
            c = ('Площадь квартиры ', str(dom_plosh(a, b)), ' м^2')
            doc.add_paragraph(c)
        else:
            msg = "Приложение обнаружило неизвестную ошибку"
            mb.showerror("Ошибка", msg)
    else:
        msg = "Приложение обнаружило неизвестную ошибку"
        mb.showerror("Ошибка", msg)


def utog2():
    y = comboExample.current()
    pop1 = shirina.get()
    pop2 = dlina.get()
    pop3 = kolvo_tt.get()
    if ((pop1.isdigit()) and (pop2.isdigit())) and ((len(pop1) != 0) and (len(pop2) != 0)):
        vivod_tepla.delete(0, 'end')
        if y == 0:
            a = int(vivod_ploshadi.get())
            c = str(comnata_energy(a))
            vivod_tepla.insert(0, c)
            c = ('Энергопотребление комнаты ', str(comnata_energy(a)), ' Вт')
            doc.add_paragraph(c)
        elif y == 1:
            a = int(vivod_ploshadi.get())
            c = str(kvartura_energy(a))
            vivod_tepla.insert(0, c)
            c = ('Энергопотребление комнаты ', str(kvartura_energy(a)), ' Вт')
            doc.add_paragraph(c)
        elif y == 2 and (len(pop3) != 0):
            d = int(kolvo_tt.get())
            a = int(vivod_ploshadi.get())
            c = str(dom_energy(a,d))
            vivod_tepla.insert(0, c)
            c = ('Энергопотребление комнаты ', str(dom_energy(a,d)), ' Вт')
            doc.add_paragraph(c)
        else:
            msg = "Приложение обнаружило неизвестную ошибку"
            mb.showerror("Ошибка", msg)
    else:
        msg = "Приложение обнаружило неизвестную ошибку"
        mb.showerror("Ошибка", msg)











#def add_label():
#    label = tk.Label(win,text='Хватит')
#    label.place(x = 100, y = 100)

win = tk.Tk()
win.title('Приложение не тупое')
win.geometry("400x300+600+200")
win.resizable(False, False)
win.config(bg='#008080')

label1_1 = tk.Label(win, text='Выбирите какое помещение нужно посчитать').place(x = 0, y = 0)

comboExample = ttk.Combobox(win,
                            values=[
                                    "Комната",
                                    "Квартира",
                                    "Дом"])
comboExample.current(0)
comboExample.place(x=260, y=0)

label1_2 = tk.Label(win, text='Введите ширину')
label1_2.place(x=0, y=30)

label1_3 = tk.Label(win, text='Введите длину')
label1_3.place(x=0, y=60)

label1_3 = tk.Label(win, text='Количество комнат, для квартиры')
label1_3.place(x=0, y=90)

label1_4 = tk.Label(win, text='Введите количество этажей, для дома')
label1_4.place(x=0, y=120)

shirina = tk.Entry(win)
shirina.place(x=100, y=30)

dlina = tk.Entry(win)
dlina.place(x=90, y=60)

vivod_ploshadi = tk.Entry(win)
vivod_ploshadi.place( x=250, y=250)

vivod_tepla = tk.Entry(win)
vivod_tepla.place( x=250, y=275)

kolvo_tt = tk.Entry(win)
kolvo_tt.place( x=220, y=122)

kolvo_com = tk.Entry(win)
kolvo_com.place( x=200, y=92)

but1 = tk.Button(win, text='Высчитать площадь', command=utog)
but1.place(x=30, y=250)

but1 = tk.Button(win, text='Высчитать потребление тепла', command=utog2)
but1.place(x=0, y=275)

win.mainloop()

# Запись в текстовый файл и вывод
doc.save('report_lab10.docx')
text = []
for paragraph in doc.paragraphs:
    text.append(paragraph.text)
print('\n'.join(text))

